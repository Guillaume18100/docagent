import os
import logging
import tempfile
from django.conf import settings
from gpt4all import GPT4All
import docx
from docx.shared import Pt
import io

# Configure logging
logger = logging.getLogger(__name__)

# Global variable to store the loaded model
GPT_MODEL = None

def load_gpt_model():
    """
    Load the GPT model for document generation
    
    Returns:
        GPT4All: Loaded model
    """
    global GPT_MODEL
    
    if GPT_MODEL is None:
        try:
            # Create models directory if it doesn't exist
            os.makedirs(settings.GPT_MODEL_PATH, exist_ok=True)
            
            # Load a GPT4All model
            # This will download the model if it doesn't exist
            model_name = "orca-mini-3b-gguf2-q4_0.gguf"  # Small model for faster loading
            GPT_MODEL = GPT4All(model_name=model_name, model_path=settings.GPT_MODEL_PATH)
            logger.info(f"Loaded GPT model: {model_name}")
        except Exception as e:
            logger.error(f"Error loading GPT model: {str(e)}")
            raise
    
    return GPT_MODEL

def generate_document_content(prompt, template=None, source_documents=None, parameters=None):
    """
    Generate document content using GPT model
    
    Args:
        prompt (str): The prompt for document generation
        template (DocumentTemplate, optional): Template to use for generation
        source_documents (list, optional): List of source documents to use as context
        parameters (dict, optional): Additional parameters for generation
        
    Returns:
        str: Generated document content
    """
    try:
        # Load the model
        model = load_gpt_model()
        
        # Prepare the context from source documents
        context = ""
        if source_documents:
            context = "Context information:\n\n"
            for doc in source_documents:
                # Limit the context to avoid token limits
                if doc.extracted_text:
                    context += f"--- From {doc.title} ---\n"
                    context += doc.extracted_text[:2000] + "...\n\n"
        
        # Prepare the template content
        template_content = ""
        if template:
            template_content = f"Template to follow:\n\n{template.template_content}\n\n"
            
            # Add template variables if available
            if template.variables:
                template_content += "Template variables:\n"
                for key, value in template.variables.items():
                    template_content += f"- {key}: {value}\n"
        
        # Prepare the final prompt
        system_prompt = (
            "You are an AI assistant specialized in document generation. "
            "Your task is to create a professional document based on the provided information. "
            "Follow any templates or guidelines provided. "
            "Use formal language and proper formatting."
        )
        
        final_prompt = f"{system_prompt}\n\n"
        
        if context:
            final_prompt += f"{context}\n\n"
            
        if template_content:
            final_prompt += f"{template_content}\n\n"
            
        final_prompt += f"Task: {prompt}\n\n"
        
        # Add any additional parameters
        if parameters:
            for key, value in parameters.items():
                if key != 'prompt':  # Skip the prompt as we already included it
                    final_prompt += f"{key}: {value}\n"
        
        # Generate the document content
        response = model.generate(
            final_prompt,
            max_tokens=2000,
            temp=0.7,
            top_k=40,
            top_p=0.4,
            repeat_penalty=1.18
        )
        
        return response
    except Exception as e:
        logger.error(f"Error generating document content: {str(e)}")
        raise

def create_docx_file(content, title):
    """
    Create a DOCX file from the generated content
    
    Args:
        content (str): The document content
        title (str): The document title
        
    Returns:
        BytesIO: DOCX file as a BytesIO object
    """
    try:
        # Create a new document
        doc = docx.Document()
        
        # Add title
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        
        # Add content
        doc.add_paragraph("\n")
        
        # Split content by paragraphs and add each paragraph
        paragraphs = content.split("\n")
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph)
        
        # Save the document to a BytesIO object
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        
        return f
    except Exception as e:
        logger.error(f"Error creating DOCX file: {str(e)}")
        raise

def generate_document(generated_document):
    """
    Generate a document based on the GeneratedDocument model
    
    Args:
        generated_document: GeneratedDocument model instance
        
    Returns:
        bool: True if generation was successful, False otherwise
    """
    try:
        # Update status
        generated_document.generation_status = 'processing'
        generated_document.save()
        
        # Get parameters
        prompt = generated_document.parameters.get('prompt', '')
        template = generated_document.template
        source_documents = generated_document.source_documents.all()
        
        # Generate content
        content = generate_document_content(
            prompt=prompt,
            template=template,
            source_documents=source_documents,
            parameters=generated_document.parameters
        )
        
        # Save the generated content
        generated_document.content = content
        
        # Create a DOCX file
        docx_file = create_docx_file(content, generated_document.title)
        
        # Save the file to the model
        file_name = f"{generated_document.title.replace(' ', '_')}.docx"
        generated_document.file.save(file_name, docx_file)
        
        # Update status
        generated_document.generation_status = 'completed'
        generated_document.save()
        
        return True
    except Exception as e:
        logger.error(f"Error generating document {generated_document.id}: {str(e)}")
        generated_document.generation_status = 'failed'
        generated_document.error_message = str(e)
        generated_document.save()
        return False 