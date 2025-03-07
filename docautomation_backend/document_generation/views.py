from django.shortcuts import render
from rest_framework import viewsets, status
from rest_framework.decorators import action, permission_classes
from rest_framework.permissions import AllowAny, IsAuthenticated
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
from django.http import FileResponse
from django.shortcuts import get_object_or_404
import os
import threading
import io
from datetime import datetime
import traceback

from .models import DocumentTemplate, GeneratedDocument
from document_processing.models import Document

# Serializers
from rest_framework import serializers

# Try to import GPT4All
GPT4ALL_AVAILABLE = False
GPT4ALL_MODEL = None

try:
    from gpt4all import GPT4All
    
    # Initialize in a separate thread to avoid blocking
    def initialize_gpt4all():
        global GPT4ALL_AVAILABLE, GPT4ALL_MODEL
        
        try:
            print("Initializing GPT4All model...")
            # Use a different model that's more reliably available
            model_name = "orca-mini-3b-gguf2-q4_0"  # Changed from ggml-gpt4all-j-v1.3-groovy
            
            try:
                # First try to initialize with the model
                GPT4ALL_MODEL = GPT4All(model_name)
                GPT4ALL_AVAILABLE = True
                print("GPT4All model initialized successfully!")
            except ValueError as e:
                if "404" in str(e):
                    # Model download failed, try a different model
                    print(f"Failed to download model {model_name}, trying a fallback model...")
                    fallback_model = "mistral-7b-instruct-v0.2.Q4_0"
                    try:
                        GPT4ALL_MODEL = GPT4All(fallback_model)
                        GPT4ALL_AVAILABLE = True
                        print(f"Fallback GPT4All model {fallback_model} initialized successfully!")
                    except Exception as e2:
                        print(f"Failed to initialize fallback model: {str(e2)}")
                        # Final fallback - disable GPT4All
                        GPT4ALL_AVAILABLE = False
                else:
                    raise
        except Exception as e:
            print(f"Error initializing GPT4All model: {str(e)}")
            traceback.print_exc()
            GPT4ALL_AVAILABLE = False
    
    # Start initialization in background
    threading.Thread(target=initialize_gpt4all).start()
    
except ImportError:
    print("GPT4All not available. Install with: pip install gpt4all")
    GPT4ALL_AVAILABLE = False

class DocumentTemplateSerializer(serializers.ModelSerializer):
    class Meta:
        model = DocumentTemplate
        fields = '__all__'
        read_only_fields = ['id', 'created_at', 'updated_at']

class GeneratedDocumentSerializer(serializers.ModelSerializer):
    class Meta:
        model = GeneratedDocument
        fields = '__all__'
        read_only_fields = ['id', 'created_at', 'updated_at']

# Document Template ViewSet
class DocumentTemplateViewSet(viewsets.ModelViewSet):
    """ViewSet for document templates"""
    queryset = DocumentTemplate.objects.all().order_by('-created_at')
    serializer_class = DocumentTemplateSerializer
    parser_classes = (MultiPartParser, FormParser, JSONParser)
    
    def get_permissions(self):
        """Return appropriate permissions"""
        return [AllowAny()]

# Define generate_document as a standalone function so it can be imported
def generate_document(generated_doc):
    """Background task to generate document content"""
    try:
        # Update status
        generated_doc.status = 'generating'
        generated_doc.save()
        
        # Get reference document texts if available
        reference_texts = []
        for doc in generated_doc.reference_documents.all():
            if doc.extracted_text:
                reference_texts.append({
                    'title': doc.title,
                    'text': doc.extracted_text
                })
        
        # Generate content using GPT4All if available
        content = ""
        if GPT4ALL_AVAILABLE and GPT4ALL_MODEL:
            content = generate_with_gpt4all(generated_doc.prompt, reference_texts, generated_doc.output_format)
        
        # If GPT4All failed or isn't available, use fallback generation
        if not content:
            content = generate_fallback_content(generated_doc.title, generated_doc.prompt, reference_texts)
        
        # Save the generated content
        generated_doc.content = content
        
        # Generate an actual file based on the requested format
        filename = f"{generated_doc.title.replace(' ', '_')}.{generated_doc.output_format}"
        
        # Create the appropriate file format
        file_content, file_path = create_document_file(
            content, 
            filename, 
            generated_doc.output_format
        )
        
        # Update the document with the file and mark as completed
        generated_doc.file.name = file_path
        generated_doc.status = 'completed'
        generated_doc.save()
        
    except Exception as e:
        import traceback
        print(f"Error generating document: {str(e)}")
        print(traceback.format_exc())
        
        # Update status to failed
        generated_doc.status = 'failed'
        generated_doc.error_message = str(e)
        generated_doc.save()

def generate_with_gpt4all(prompt, reference_texts, output_format):
    """Generate document content using GPT4All"""
    try:
        # Prepare the system prompt
        system_prompt = "You are an AI assistant that generates professional documents based on user requirements."
        
        # Prepare the user prompt with reference documents
        user_prompt = f"Please generate a {output_format.upper()} document based on the following requirements:\n\n{prompt}\n\n"
        
        # Add reference documents if available
        if reference_texts:
            user_prompt += "Reference documents:\n\n"
            for i, doc in enumerate(reference_texts):
                user_prompt += f"Document {i+1}: {doc['title']}\n"
                # Limit text length to avoid context window issues
                user_prompt += f"{doc['text'][:1000]}...\n\n"
        
        # Add format instructions
        if output_format == 'markdown':
            user_prompt += "\nPlease format the document in Markdown syntax."
        elif output_format == 'html':
            user_prompt += "\nPlease format the document in HTML syntax."
        
        # Generate with GPT4All
        response = GPT4ALL_MODEL.generate(
            user_prompt,
            max_tokens=2000,
            temp=0.7,
            top_k=40,
            top_p=0.9,
            repeat_penalty=1.18,
            repeat_last_n=64,
            n_batch=8,
            n_predict=None,
            streaming=False,
        )
        
        return response
        
    except Exception as e:
        print(f"Error generating with GPT4All: {str(e)}")
        traceback.print_exc()
        return ""

def generate_fallback_content(title, prompt, reference_texts):
    """Generate fallback content when GPT4All is not available"""
    content = f"""# {title}

## Generated based on your prompt:
"{prompt}"

## Document Content
This is a generated document based on your requirements. In a production environment, 
this would be generated by GPT4All or another AI model.

"""
    
    # Add sections based on common document structures
    content += """
## Executive Summary
This document addresses the requirements specified in your prompt.

## Introduction
This section would introduce the main topics and purpose of the document.

## Main Content
This section would contain the primary content generated based on your requirements.
"""
    
    # Add reference document information if available
    if reference_texts:
        content += "\n## Reference Documents\n"
        for i, doc in enumerate(reference_texts):
            content += f"\n### Document {i+1}: {doc['title']}\n"
            content += f"Excerpt: {doc['text'][:200]}...\n"
    else:
        content += "\n## References\nNo reference documents were provided."
    
    # Add conclusion
    content += """
## Conclusion
This concludes the generated document based on your requirements.
"""
    
    return content

def create_document_file(content, filename, output_format):
    """Create a file in the requested format"""
    from django.core.files.base import ContentFile
    from django.core.files.storage import default_storage
    
    try:
        if output_format == 'docx':
            # Create a Word document
            try:
                from docx import Document
                doc = Document()
                
                # Split content by lines and add to document
                lines = content.split('\n')
                for line in lines:
                    if line.startswith('# '):
                        # Main heading
                        doc.add_heading(line[2:], level=0)
                    elif line.startswith('## '):
                        # Subheading
                        doc.add_heading(line[3:], level=1)
                    elif line.startswith('### '):
                        # Sub-subheading
                        doc.add_heading(line[4:], level=2)
                    elif line.strip():
                        # Regular paragraph
                        doc.add_paragraph(line)
                
                # Save to memory
                file_stream = io.BytesIO()
                doc.save(file_stream)
                file_stream.seek(0)
                
                # Save to storage
                file_path = default_storage.save(
                    f"generated_documents/{filename}", 
                    ContentFile(file_stream.read())
                )
                return content, file_path
                
            except ImportError:
                print("python-docx not available. Falling back to text file.")
                # Fall back to text file
                file_content = content.encode('utf-8')
                file_path = default_storage.save(
                    f"generated_documents/{filename}", 
                    ContentFile(file_content)
                )
                return content, file_path
                
        elif output_format == 'pdf':
            # Create a PDF document
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet
                
                # Create PDF in memory
                file_stream = io.BytesIO()
                doc = SimpleDocTemplate(file_stream, pagesize=letter)
                styles = getSampleStyleSheet()
                
                # Process content
                story = []
                lines = content.split('\n')
                for line in lines:
                    if line.startswith('# '):
                        # Main heading
                        story.append(Paragraph(line[2:], styles['Title']))
                        story.append(Spacer(1, 12))
                    elif line.startswith('## '):
                        # Subheading
                        story.append(Paragraph(line[3:], styles['Heading1']))
                        story.append(Spacer(1, 10))
                    elif line.startswith('### '):
                        # Sub-subheading
                        story.append(Paragraph(line[4:], styles['Heading2']))
                        story.append(Spacer(1, 8))
                    elif line.strip():
                        # Regular paragraph
                        story.append(Paragraph(line, styles['Normal']))
                        story.append(Spacer(1, 6))
                
                # Build PDF
                doc.build(story)
                file_stream.seek(0)
                
                # Save to storage
                file_path = default_storage.save(
                    f"generated_documents/{filename}", 
                    ContentFile(file_stream.read())
                )
                return content, file_path
                
            except ImportError:
                print("ReportLab not available. Falling back to text file.")
                # Fall back to text file
                file_content = content.encode('utf-8')
                file_path = default_storage.save(
                    f"generated_documents/{filename}", 
                    ContentFile(file_content)
                )
                return content, file_path
        
        else:
            # For other formats (txt, markdown, html), just save as is
            file_content = content.encode('utf-8')
            file_path = default_storage.save(
                f"generated_documents/{filename}", 
                ContentFile(file_content)
            )
            return content, file_path
            
    except Exception as e:
        print(f"Error creating document file: {str(e)}")
        traceback.print_exc()
        
        # Fall back to simple text file
        file_content = content.encode('utf-8')
        file_path = default_storage.save(
            f"generated_documents/{filename.split('.')[0]}.txt", 
            ContentFile(file_content)
        )
        return content, file_path

# Generated Document ViewSet
class GeneratedDocumentViewSet(viewsets.ModelViewSet):
    """ViewSet for generated documents"""
    queryset = GeneratedDocument.objects.all().order_by('-created_at')
    serializer_class = GeneratedDocumentSerializer
    parser_classes = (MultiPartParser, FormParser, JSONParser)
    
    def get_permissions(self):
        """Return appropriate permissions"""
        return [AllowAny()]
    
    def create(self, request, *args, **kwargs):
        """Create a new document generation request"""
        try:
            # Extract data from request
            title = request.data.get('title', f"Generated Document {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            prompt = request.data.get('prompt', '')
            output_format = request.data.get('output_format', 'docx')
            template_id = request.data.get('template_id')
            document_ids = request.data.get('document_ids', [])
            
            if not prompt:
                return Response(
                    {'error': 'Prompt is required'},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # Create the generated document
            generated_doc = GeneratedDocument(
                title=title,
                prompt=prompt,
                output_format=output_format,
                status='pending'
            )
            
            # Add template if provided
            if template_id:
                try:
                    template = DocumentTemplate.objects.get(id=template_id)
                    generated_doc.template = template
                except DocumentTemplate.DoesNotExist:
                    pass
            
            generated_doc.save()
            
            # Add reference documents if provided
            if document_ids and isinstance(document_ids, list):
                for doc_id in document_ids:
                    try:
                        doc = Document.objects.get(id=doc_id)
                        generated_doc.reference_documents.add(doc)
                    except Document.DoesNotExist:
                        pass
            
            # Start generation in background thread
            threading.Thread(target=generate_document, args=(generated_doc,)).start()
            
            # Return the created document
            serializer = self.get_serializer(generated_doc)
            headers = self.get_success_headers(serializer.data)
            return Response(
                serializer.data, 
                status=status.HTTP_201_CREATED, 
                headers=headers
            )
            
        except Exception as e:
            import traceback
            print(f"Error in create: {str(e)}")
            print(traceback.format_exc())
            return Response(
                {"error": str(e)}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    @action(detail=True, methods=['post'])
    def regenerate(self, request, pk=None):
        """Regenerate a document with the same parameters"""
        generated_doc = self.get_object()
        
        # Start regeneration in background thread
        threading.Thread(target=generate_document, args=(generated_doc,)).start()
        
        return Response(
            {'status': 'Document regeneration started'},
            status=status.HTTP_202_ACCEPTED
        )
    
    @action(detail=True, methods=['get'])
    def download(self, request, pk=None):
        """Download the generated document file"""
        generated_doc = self.get_object()
        
        if not generated_doc.file:
            return Response(
                {'error': 'No file available for this document'},
                status=status.HTTP_404_NOT_FOUND
            )
        
        try:
            file_path = generated_doc.file.path
            
            if not os.path.exists(file_path):
                return Response(
                    {'error': 'File not found on server'},
                    status=status.HTTP_404_NOT_FOUND
                )
            
            # Determine content type
            content_type = 'application/octet-stream'  # Default
            if generated_doc.output_format == 'pdf':
                content_type = 'application/pdf'
            elif generated_doc.output_format == 'docx':
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            elif generated_doc.output_format == 'txt' or generated_doc.output_format == 'markdown':
                content_type = 'text/plain'
            elif generated_doc.output_format == 'html':
                content_type = 'text/html'
            
            # Serve the file
            response = FileResponse(
                open(file_path, 'rb'),
                content_type=content_type
            )
            
            # Set filename
            filename = f"{generated_doc.title}.{generated_doc.output_format}"
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            
            return response
            
        except Exception as e:
            import traceback
            print(f"Error in download: {str(e)}")
            print(traceback.format_exc())
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
