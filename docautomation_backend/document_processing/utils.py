import os
import logging
import threading

# Configure logging
logger = logging.getLogger(__name__)

# Try to import optional dependencies
try:
    import pytesseract
    from PIL import Image
    HAVE_TESSERACT = True
except ImportError:
    logger.warning("pytesseract or PIL not available. OCR features will be disabled.")
    HAVE_TESSERACT = False

try:
    import tika
    from tika import parser
    # Initialize tika
    tika.initVM()
    HAVE_TIKA = True
except ImportError:
    logger.warning("tika-python not available. Full-text extraction from PDFs and Office documents will be limited.")
    HAVE_TIKA = False

try:
    from pdf2image import convert_from_path
    HAVE_PDF2IMAGE = True
except ImportError:
    logger.warning("pdf2image not available. PDF OCR capabilities will be limited.")
    HAVE_PDF2IMAGE = False

try:
    import docx
    HAVE_DOCX = True
except ImportError:
    logger.warning("python-docx not available. DOCX extraction capabilities will be limited.")
    HAVE_DOCX = False


def extract_text_from_pdf(file_path):
    """
    Extract text from PDF using Tika or fallback methods
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        tuple: (extracted_text, metadata)
    """
    try:
        if HAVE_TIKA:
            parsed = parser.from_file(file_path)
            text = parsed.get('content', '')
            metadata = parsed.get('metadata', {})
            
            # If text extraction failed or returned very little text, try OCR
            if not text or len(text.strip()) < 100:
                logger.info(f"Tika extracted minimal text from {file_path}, trying OCR")
                if HAVE_TESSERACT and HAVE_PDF2IMAGE:
                    return extract_text_from_pdf_with_ocr(file_path)
                else:
                    logger.warning("OCR fallback not available - missing dependencies")
                
            return text, metadata
        else:
            # Simple fallback if tika is not available
            logger.warning("Using simple text extraction - Tika not available")
            return f"Text extraction not available for {os.path.basename(file_path)}", {}
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
        return "", {}

def extract_text_from_pdf_with_ocr(file_path):
    """
    Extract text from PDF using OCR when Tika fails
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        tuple: (extracted_text, metadata)
    """
    if not (HAVE_TESSERACT and HAVE_PDF2IMAGE):
        return "OCR processing not available - missing dependencies", {}
        
    try:
        # Convert PDF to images
        images = convert_from_path(file_path)
        
        # Extract text from each image
        text = ""
        for i, image in enumerate(images):
            temp_img_path = f"/tmp/page_{i}.png"
            image.save(temp_img_path, 'PNG')
            
            # Extract text using pytesseract
            page_text = pytesseract.image_to_string(Image.open(temp_img_path))
            text += f"\n\n--- Page {i+1} ---\n\n{page_text}"
            
            # Clean up temporary file
            os.remove(temp_img_path)
            
        return text, {"ocr_processed": True}
    except Exception as e:
        logger.error(f"Error extracting text from PDF with OCR {file_path}: {str(e)}")
        return "", {"ocr_error": str(e)}

def extract_text_from_docx(file_path):
    """
    Extract text from DOCX file
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        tuple: (extracted_text, metadata)
    """
    try:
        if HAVE_TIKA:
            # Try with Tika first
            parsed = parser.from_file(file_path)
            text = parsed.get('content', '')
            metadata = parsed.get('metadata', {})
            
            # If Tika fails, use python-docx as fallback
            if not text and HAVE_DOCX:
                logger.info(f"Tika failed to extract text from {file_path}, using python-docx")
                doc = docx.Document(file_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                metadata = {"fallback_extraction": "python-docx"}
        elif HAVE_DOCX:
            # Use python-docx directly
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            metadata = {"extraction_method": "python-docx"}
        else:
            # No extraction methods available
            return f"Text extraction not available for {os.path.basename(file_path)}", {}
            
        return text, metadata
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
        return "", {}

def extract_text_from_image(file_path):
    """
    Extract text from image using OCR
    
    Args:
        file_path (str): Path to the image file
        
    Returns:
        tuple: (extracted_text, metadata)
    """
    if not HAVE_TESSERACT:
        return "OCR processing not available - pytesseract not installed", {}
        
    try:
        # Use pytesseract for OCR
        text = pytesseract.image_to_string(Image.open(file_path))
        metadata = {"ocr_engine": "pytesseract"}
        return text, metadata
    except Exception as e:
        logger.error(f"Error extracting text from image {file_path}: {str(e)}")
        return "", {}

def extract_text_from_txt(file_path):
    """
    Extract text from plain text file
    
    Args:
        file_path (str): Path to the text file
        
    Returns:
        tuple: (extracted_text, metadata)
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        return text, {"file_type": "text/plain"}
    except UnicodeDecodeError:
        # Try with different encoding if UTF-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
            return text, {"file_type": "text/plain", "encoding": "latin-1"}
        except Exception as e:
            logger.error(f"Error reading text file with latin-1 encoding {file_path}: {str(e)}")
            return "", {}
    except Exception as e:
        logger.error(f"Error reading text file {file_path}: {str(e)}")
        return "", {}

def process_document(document):
    """
    Process a document based on its type
    
    Args:
        document: Document model instance
        
    Returns:
        bool: True if processing was successful, False otherwise
    """
    try:
        logger.info(f"Processing document: {document.title}")
        
        file_path = document.file.path
        document.processing_status = 'processing'
        document.save()
        
        # For simplicity in the prototype version, just use a basic method
        # that doesn't require external dependencies
        with open(file_path, 'rb') as f:
            # Try to read as text first
            try:
                content = f.read().decode('utf-8')
                text = content
                metadata = {"simple_extraction": True}
            except UnicodeDecodeError:
                # If it's not text, provide a placeholder
                text = f"[Processed content for {document.title}]"
                metadata = {"binary_file": True}
                
                # Try more advanced extraction if available
                if document.document_type == 'pdf' and (HAVE_TIKA or HAVE_TESSERACT):
                    text, metadata = extract_text_from_pdf(file_path)
                elif document.document_type == 'docx' and (HAVE_TIKA or HAVE_DOCX):
                    text, metadata = extract_text_from_docx(file_path)
                elif document.document_type == 'image' and HAVE_TESSERACT:
                    text, metadata = extract_text_from_image(file_path)
        
        # Update document with extracted text and metadata
        document.extracted_text = text
        document.metadata = metadata
        document.processing_status = 'completed'
        document.save()
        
        logger.info(f"Document {document.id} processed successfully")
        return True
    except Exception as e:
        logger.error(f"Error processing document {document.id}: {str(e)}")
        document.processing_status = 'failed'
        document.error_message = str(e)
        document.save()
        return False 